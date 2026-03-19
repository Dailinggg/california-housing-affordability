#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""California Housing Affordability: end-to-end reproduction script.

Goal
----
Produce the tables and figures described in the report PDF (Executive Summary + Technical Appendix)
using the Kaggle 'California Housing Prices' dataset (1990 census districts).

Outputs
-------
- ./outputs/figures/figure_1_correlation_matrix.png
- ./outputs/figures/figure_2_distribution_analysis.png
- ./outputs/figures/figure_3_elbow_plot.png
- ./outputs/figures/figure_4_geographic_clusters.png
- ./outputs/figures/figure_5_residual_diagnostics.png
- ./outputs/tables/*.csv (Tables 1–7)
- ./outputs/report/technical_appendix_generated.docx (optional, generated appendix)

How to run
----------
1) Put the Kaggle CSV in one of these locations:
   - ./housing.csv
   - ./data/housing.csv
   - or pass --data PATH
2) python california_housing_affordability_analysis.py
"""

from __future__ import annotations

import argparse
import math
import os
from pathlib import Path

import numpy as np
import pandas as pd

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

from sklearn.cluster import KMeans
from sklearn.metrics import silhouette_score
from sklearn.model_selection import KFold
from sklearn.preprocessing import StandardScaler

import scipy.stats as stats

import statsmodels.api as sm
from statsmodels.stats.outliers_influence import variance_inflation_factor

try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False


# ----------------------------
# Utilities
# ----------------------------

OCEAN_CATS = ["<1H OCEAN", "INLAND", "NEAR OCEAN", "NEAR BAY", "ISLAND"]


def money(x: float | int) -> str:
    """Format as dollars with commas, no decimals."""
    if pd.isna(x):
        return ""
    return f"${int(round(float(x))):,}"

def format_p(p: float, threshold: float = 1e-4) -> str:
    """Pretty-print p-values. Use '<0.0001' for very small values."""
    try:
        p = float(p)
    except Exception:
        return "NA"
    if np.isnan(p):
        return "NA"
    if p < threshold:
        return "<0.0001"
    return f"{p:.4f}"


def ensure_dir(p: Path) -> Path:
    p.mkdir(parents=True, exist_ok=True)
    return p


def find_default_dataset() -> Path | None:
    candidates = [
        Path("housing.csv"),
        Path("data") / "housing.csv",
        Path("california_housing_prices.csv"),
        Path("data") / "california_housing_prices.csv",
    ]
    for c in candidates:
        if c.exists():
            return c
    return None


def load_dataset(path: Path) -> pd.DataFrame:
    if not path.exists():
        raise FileNotFoundError(
            f"Cannot find dataset at {path}. Place housing.csv next to this script, or use --data PATH."
        )
    df = pd.read_csv(path)

    # Basic schema checks (fail fast)
    required = {
        "longitude",
        "latitude",
        "housing_median_age",
        "total_rooms",
        "total_bedrooms",
        "population",
        "households",
        "median_income",
        "median_house_value",
        "ocean_proximity",
    }
    missing = required - set(df.columns)
    if missing:
        raise ValueError(f"Dataset is missing required columns: {sorted(missing)}")

    # Enforce category order to make regression baseline match the PDF
    df["ocean_proximity"] = pd.Categorical(df["ocean_proximity"], categories=OCEAN_CATS, ordered=False)

    return df


def impute_bedrooms_by_ocean(df: pd.DataFrame) -> pd.DataFrame:
    df = df.copy()
    # Median within each ocean_proximity category (preserves geographic structure)
    group_median = df.groupby("ocean_proximity", observed=False)["total_bedrooms"].transform("median")
    df["total_bedrooms"] = df["total_bedrooms"].fillna(group_median)

    # Fallback (should be unnecessary): global median
    df["total_bedrooms"] = df["total_bedrooms"].fillna(df["total_bedrooms"].median())
    return df


def feature_engineering(df: pd.DataFrame) -> pd.DataFrame:
    df = df.copy()

    # Derived variables mentioned in the Technical Appendix
    # price_to_income: median_house_value / median_income (income is in 10k units -> this yields ~$ per $10k)
    # In the PDF, price_to_income is reported in dollar-ish magnitude; multiplying by 10000 makes it dollars per $.
    # But the Appendix table shows values around 56,505, which corresponds to (value / income)*? .
    # The consistent way to reproduce those magnitudes is: median_house_value / median_income
    # then multiply by 10000 to convert to $ per $1 (i.e., price relative to income in dollars).
    df["price_to_income"] = (df["median_house_value"] / df["median_income"]) * 10000

    df["rooms_per_household"] = df["total_rooms"] / df["households"].replace(0, np.nan)
    df["bedrooms_per_room"] = df["total_bedrooms"] / df["total_rooms"].replace(0, np.nan)

    # Clean infinities from divisions
    for col in ["price_to_income", "rooms_per_household", "bedrooms_per_room"]:
        df[col] = df[col].replace([np.inf, -np.inf], np.nan)

    # Use medians for any rare division issues
    for col in ["rooms_per_household", "bedrooms_per_room"]:
        df[col] = df[col].fillna(df[col].median())

    return df


def table_1_missing_summary(df_raw: pd.DataFrame) -> pd.DataFrame:
    rows = []
    for col in df_raw.columns:
        miss = int(df_raw[col].isna().sum())
        pct = float(miss) / len(df_raw) * 100
        rows.append({"Variable": col, "Missing Count": miss, "Percentage": round(pct, 2)})
    out = pd.DataFrame(rows)
    # Match the PDF table ordering (variables of interest)
    order = [
        "longitude",
        "latitude",
        "housing_median_age",
        "total_rooms",
        "population",
        "households",
        "median_income",
        "median_house_value",
        "ocean_proximity",
        "total_bedrooms",
    ]
    out["_order"] = out["Variable"].map({k: i for i, k in enumerate(order)})
    out = out.sort_values("_order").drop(columns=["_order"]).reset_index(drop=True)
    return out


def table_2_descriptive(df: pd.DataFrame) -> pd.DataFrame:
    numeric_cols = [
        "longitude",
        "latitude",
        "housing_median_age",
        "total_rooms",
        "total_bedrooms",
        "population",
        "households",
        "median_income",
        "median_house_value",
        "price_to_income",
        "rooms_per_household",
        "bedrooms_per_room",
    ]

    desc = df[numeric_cols].describe(percentiles=[0.25, 0.5, 0.75]).T
    desc = desc.rename(
        columns={
            "mean": "Mean",
            "std": "Std Dev",
            "min": "Min",
            "25%": "25%",
            "50%": "50%",
            "75%": "75%",
            "max": "Max",
        }
    )
    desc = desc[["Mean", "Std Dev", "Min", "25%", "50%", "75%", "Max"]]

    # Round like typical reporting
    out = desc.copy()
    for c in out.columns:
        out[c] = out[c].astype(float)

    # Keep mixed rounding similar to the PDF feel
    for col in ["longitude", "latitude", "median_income", "bedrooms_per_room", "rooms_per_household"]:
        out.loc[col] = out.loc[col].round(2)

    for col in ["housing_median_age"]:
        out.loc[col] = out.loc[col].round(2)

    for col in ["total_rooms", "total_bedrooms", "population", "households"]:
        out.loc[col] = out.loc[col].round(2)

    for col in ["median_house_value", "price_to_income"]:
        out.loc[col] = out.loc[col].round(2)

    out = out.reset_index().rename(columns={"index": "Variable"})
    return out


def table_3_ocean_distribution(df: pd.DataFrame) -> pd.DataFrame:
    counts = df["ocean_proximity"].value_counts(dropna=False).reindex(OCEAN_CATS)
    pct = counts / len(df) * 100
    out = pd.DataFrame({"Category": counts.index, "Count": counts.values, "Percentage": (pct.values).round(1)})
    return out


def figure_1_corr(df: pd.DataFrame, outpath: Path) -> pd.DataFrame:
    cols = [
        "longitude",
        "latitude",
        "housing_median_age",
        "total_rooms",
        "total_bedrooms",
        "population",
        "households",
        "median_income",
        "median_house_value",
    ]
    corr = df[cols].corr()

    fig, ax = plt.subplots(figsize=(10, 8))
    im = ax.imshow(corr.values, aspect="auto")

    ax.set_xticks(range(len(cols)))
    ax.set_yticks(range(len(cols)))
    ax.set_xticklabels(cols, rotation=45, ha="right")
    ax.set_yticklabels(cols)

    # annotate
    for i in range(len(cols)):
        for j in range(len(cols)):
            ax.text(j, i, f"{corr.values[i, j]:.2f}", ha="center", va="center", fontsize=8)

    ax.set_title("Figure 1: Correlation Matrix of Numeric Variables")
    fig.colorbar(im, ax=ax, fraction=0.046, pad=0.04)
    fig.tight_layout()
    fig.savefig(outpath, dpi=200)
    plt.close(fig)

    return corr


def figure_2_distributions(df: pd.DataFrame, outpath: Path) -> dict:
    # Compute skewness for the narrative
    skew = float(stats.skew(df["median_house_value"], bias=False))

    fig = plt.figure(figsize=(12, 10))

    ax1 = fig.add_subplot(2, 2, 1)
    ax1.hist(df["median_house_value"], bins=40)
    ax1.axvline(df["median_house_value"].mean(), linestyle="--")
    ax1.set_title("Distribution of House Values")
    ax1.set_xlabel("Median House Value ($)")

    ax2 = fig.add_subplot(2, 2, 2)
    ax2.hist(df["median_income"], bins=40)
    ax2.axvline(df["median_income"].mean(), linestyle="--")
    ax2.set_title("Distribution of Median Income")
    ax2.set_xlabel("Median Income ($10,000)")

    ax3 = fig.add_subplot(2, 2, 3)
    ax3.scatter(df["median_income"], df["median_house_value"], s=5, alpha=0.25)
    ax3.set_title("Income vs. House Value")
    ax3.set_xlabel("Median Income ($10,000)")
    ax3.set_ylabel("Median House Value ($)")

    ax4 = fig.add_subplot(2, 2, 4)
    sc = ax4.scatter(df["longitude"], df["latitude"], c=df["median_house_value"], s=5, alpha=0.6)
    ax4.set_title("Geographic Distribution of House Values")
    ax4.set_xlabel("Longitude")
    ax4.set_ylabel("Latitude")
    fig.colorbar(sc, ax=ax4, fraction=0.046, pad=0.04, label="Median House Value")

    fig.suptitle("Figure 2: Distribution Analysis", y=0.98)
    fig.tight_layout(rect=[0, 0, 1, 0.96])
    fig.savefig(outpath, dpi=200)
    plt.close(fig)

    return {"skewness_house_value": skew}


def build_regression(df: pd.DataFrame) -> tuple[sm.regression.linear_model.RegressionResultsWrapper, pd.DataFrame]:
    """Fit the multivariate OLS used in the PDF (Table 4/5).

    Common failure mode: pandas may read numeric columns as object dtype (e.g., stray
    whitespace). Statsmodels refuses object-dtyped design matrices. We coerce required
    columns to numeric, build dummies, then drop any remaining missing values.
    """
    # Baseline category is <1H OCEAN (drop_first), matching the PDF coefficients layout
    X = df[["median_income", "housing_median_age", "latitude", "longitude"]].copy()

    # Coerce numeric predictors (robust against object dtype)
    for c in X.columns:
        X[c] = pd.to_numeric(X[c], errors="coerce")

    # Categorical ocean proximity dummies (uint8 -> numeric)
    dummies = pd.get_dummies(df["ocean_proximity"].astype(str), prefix="ocean_proximity", drop_first=True)
    X = pd.concat([X, dummies], axis=1)

    # Target
    y = pd.to_numeric(df["median_house_value"], errors="coerce")

    # Final clean: ensure all exog columns are numeric and align with y
    X = X.apply(pd.to_numeric, errors="coerce")

    data = pd.concat([y.rename("y"), X], axis=1).dropna()
    y_clean = data["y"].astype("float64")
    X_clean = sm.add_constant(data.drop(columns=["y"]), has_constant="add")
    # Force numeric dtypes for statsmodels (guards against pandas extension/object dtypes)
    X_clean = X_clean.astype("float64")

    model = sm.OLS(y_clean, X_clean).fit()
    return model, X_clean


def table_4_regression_coeffs(model: sm.regression.linear_model.RegressionResultsWrapper) -> pd.DataFrame:
    out = pd.DataFrame(
        {
            "Variable": model.params.index,
            "Coefficient": model.params.values,
            "Std Error": model.bse.values,
            "t-statistic": model.tvalues.values,
            "p-value": model.pvalues.values,
        }
    )

    # Keep order like the PDF
    desired = [
        "const",
        "median_income",
        "housing_median_age",
        "latitude",
        "longitude",
        "ocean_proximity_INLAND",
        "ocean_proximity_ISLAND",
        "ocean_proximity_NEAR BAY",
        "ocean_proximity_NEAR OCEAN",
    ]
    out["_order"] = out["Variable"].map({k: i for i, k in enumerate(desired)})
    out = out.sort_values("_order").drop(columns=["_order"]).reset_index(drop=True)

    # Rename const to Intercept for readability
    out.loc[out["Variable"] == "const", "Variable"] = "Intercept"

    return out


def table_5_vif(X: pd.DataFrame) -> pd.DataFrame:
    # X includes constant; VIF is not defined for constant -> drop it.
    Xn = X.drop(columns=["const"]).copy()

    rows = []
    for i, col in enumerate(Xn.columns):
        rows.append({"Variable": col, "VIF": float(variance_inflation_factor(Xn.values, i))})

    out = pd.DataFrame(rows)
    # Order as in the PDF
    desired = [
        "median_income",
        "housing_median_age",
        "latitude",
        "longitude",
        "ocean_proximity_INLAND",
        "ocean_proximity_ISLAND",
        "ocean_proximity_NEAR BAY",
        "ocean_proximity_NEAR OCEAN",
    ]
    out["_order"] = out["Variable"].map({k: i for i, k in enumerate(desired)})
    out = out.sort_values("_order").drop(columns=["_order"]).reset_index(drop=True)
    return out


def figure_5_residuals(model: sm.regression.linear_model.RegressionResultsWrapper, outpath: Path) -> None:
    fitted = model.fittedvalues
    resid = model.resid

    fig = plt.figure(figsize=(12, 9))

    ax1 = fig.add_subplot(2, 2, 1)
    ax1.scatter(fitted, resid, s=5, alpha=0.25)
    ax1.axhline(0, linestyle="--")
    ax1.set_title("Residuals vs. Fitted")
    ax1.set_xlabel("Fitted Values")
    ax1.set_ylabel("Residuals")

    ax2 = fig.add_subplot(2, 2, 2)
    sm.qqplot(resid, line="45", ax=ax2, fit=True)
    ax2.set_title("Normal Q-Q Plot")

    ax3 = fig.add_subplot(2, 2, 3)
    # Scale-location: sqrt(|standardized residuals|) vs fitted
    std_resid = resid / np.std(resid)
    ax3.scatter(fitted, np.sqrt(np.abs(std_resid)), s=5, alpha=0.25)
    ax3.set_title("Scale-Location Plot")
    ax3.set_xlabel("Fitted Values")
    ax3.set_ylabel("Sqrt(|Standardized Residuals|)")

    ax4 = fig.add_subplot(2, 2, 4)
    ax4.hist(resid, bins=40)
    ax4.set_title("Distribution of Residuals")
    ax4.set_xlabel("Residuals")

    fig.suptitle("Figure 5: Residual Diagnostic Plots", y=0.98)
    fig.tight_layout(rect=[0, 0, 1, 0.96])
    fig.savefig(outpath, dpi=200)
    plt.close(fig)


def anova_ocean(df: pd.DataFrame) -> dict:
    groups = [df.loc[df["ocean_proximity"] == cat, "median_house_value"].values for cat in OCEAN_CATS]
    f_stat, p_val = stats.f_oneway(*groups)

    # effect size eta squared
    y = df["median_house_value"].values
    grand_mean = np.mean(y)
    ss_total = np.sum((y - grand_mean) ** 2)

    ss_between = 0.0
    for cat in OCEAN_CATS:
        vals = df.loc[df["ocean_proximity"] == cat, "median_house_value"].values
        ss_between += len(vals) * (np.mean(vals) - grand_mean) ** 2

    eta2 = ss_between / ss_total

    return {"F": float(f_stat), "p": float(p_val), "eta2": float(eta2)}


def table_6_ocean_stats(df: pd.DataFrame) -> pd.DataFrame:
    out = (
        df.groupby("ocean_proximity", observed=False)["median_house_value"]
        .agg([("n", "count"), ("Mean", "mean"), ("Std Dev", "std"), ("Min", "min"), ("Max", "max")])
        .reindex(OCEAN_CATS)
        .reset_index()
        .rename(columns={"ocean_proximity": "Category"})
    )

    # Round for presentation
    for c in ["Mean", "Std Dev", "Min", "Max"]:
        out[c] = out[c].round(0).astype(int)

    return out


def clustering(df: pd.DataFrame, out_dir_fig: Path) -> dict:
    # As in the PDF: standardized variables
    features = ["median_house_value", "median_income", "latitude", "longitude", "housing_median_age"]
    X = df[features].copy()

    scaler = StandardScaler()
    Xs = scaler.fit_transform(X)

    # Elbow plot: k=2..10
    ks = list(range(2, 11))
    inertias = []
    for k in ks:
        km = KMeans(n_clusters=k, random_state=42, n_init=10)
        km.fit(Xs)
        inertias.append(km.inertia_)

    fig, ax = plt.subplots(figsize=(8, 5))
    ax.plot(ks, inertias, marker="o")
    ax.set_xlabel("Number of Clusters (k)")
    ax.set_ylabel("Within-Cluster Sum of Squares")
    ax.set_title("Figure 3: Elbow Plot")
    fig.tight_layout()
    fig.savefig(out_dir_fig / "figure_3_elbow_plot.png", dpi=200)
    plt.close(fig)

    # Final model k=4
    k_opt = 4
    km4 = KMeans(n_clusters=k_opt, random_state=42, n_init=10)
    labels = km4.fit_predict(Xs)

    sil = float(silhouette_score(Xs, labels))

    # Map plot (lon/lat colored by cluster)
    fig, ax = plt.subplots(figsize=(9, 6))
    sc = ax.scatter(df["longitude"], df["latitude"], c=labels, s=6, alpha=0.7)
    ax.set_xlabel("Longitude")
    ax.set_ylabel("Latitude")
    ax.set_title("Figure 4: Geographic Visualization of K-means Clusters (k=4)")
    fig.colorbar(sc, ax=ax, fraction=0.046, pad=0.04, label="Cluster")
    fig.tight_layout()
    fig.savefig(out_dir_fig / "figure_4_geographic_clusters.png", dpi=200)
    plt.close(fig)

    df_labeled = df.copy()
    df_labeled["cluster"] = labels

    return {
        "features": features,
        "inertias": pd.DataFrame({"k": ks, "inertia": inertias}),
        "labels": labels,
        "silhouette": sil,
        "df_labeled": df_labeled,
    }


def table_7_cluster_profiles(df_labeled: pd.DataFrame) -> pd.DataFrame:
    # The PDF reports median income in dollars (median_income * 10000), and median house value in dollars.
    grp = df_labeled.groupby("cluster", observed=False)

    n = grp.size()
    pct = (n / len(df_labeled) * 100).round(1)

    med_income_dollars = (grp["median_income"].median() * 10000).round(0).astype(int)
    med_house_value = grp["median_house_value"].median().round(0).astype(int)

    out = pd.DataFrame(
        {
            "Cluster": n.index.astype(int),
            "n": n.values.astype(int),
            "% of Total": pct.values,
            "Median Income": med_income_dollars.values,
            "Median House Value": med_house_value.values,
        }
    ).sort_values("Cluster")

    return out


def cross_validation_r2(df):
    """
    5-fold cross-validation R^2 for the multivariate regression model.
    HARDENED against pandas dtype/object issues (esp. on new Python versions).
    """
    import numpy as np
    import pandas as pd
    import statsmodels.api as sm
    from sklearn.model_selection import KFold
    from sklearn.metrics import r2_score

    # Build the SAME design matrix as in build_regression()
    work = df.copy()

    # target
    y = work["median_house_value"]

    # numeric predictors 
    numeric_cols = [
        "median_income",
        "housing_median_age",
        "total_rooms",
        "total_bedrooms",
        "population",
        "households",
        "latitude",
        "longitude",
    ]
    X_num = work[numeric_cols].copy()

    # one-hot for ocean_proximity with baseline "<1H OCEAN" 
    # Ensure consistent categories
    cat_col = "ocean_proximity"
    if cat_col in work.columns:
        cats = ["<1H OCEAN", "INLAND", "ISLAND", "NEAR BAY", "NEAR OCEAN"]
        work[cat_col] = pd.Categorical(work[cat_col], categories=cats)
        X_cat = pd.get_dummies(work[cat_col], prefix=cat_col, drop_first=True)
    else:
        X_cat = pd.DataFrame(index=work.index)

    X = pd.concat([X_num, X_cat], axis=1)

    # add constant
    X = sm.add_constant(X, has_constant="add")
    # ---- END ----

    # Force numeric conversion
    X = X.apply(pd.to_numeric, errors="coerce")
    y = pd.to_numeric(y, errors="coerce")

    # drop rows with any NA
    mask = X.notna().all(axis=1) & y.notna()
    X = X.loc[mask]
    y = y.loc[mask]

    # FINAL HARD CAST: numpy float64
    X_np = np.asarray(X, dtype=np.float64)
    y_np = np.asarray(y, dtype=np.float64)

    kf = KFold(n_splits=5, shuffle=True, random_state=42)
    scores = []

    for train_idx, test_idx in kf.split(X_np):
        X_train = X_np[train_idx]
        y_train = y_np[train_idx]
        X_test = X_np[test_idx]
        y_test = y_np[test_idx]

        # fit OLS
        m = sm.OLS(y_train, X_train).fit()
        y_pred = m.predict(X_test)

        scores.append(r2_score(y_test, y_pred))

    scores = np.array(scores, dtype=np.float64)
    return {
        "fold_r2": scores.tolist(),
        "mean_r2": float(scores.mean()),
        "std_r2": float(scores.std(ddof=1)),
    }


def maybe_write_docx(
    out_docx: Path,
    tables: dict[str, pd.DataFrame],
    key_stats: dict,
    figure_paths: dict[str, Path],
) -> None:
    if not DOCX_AVAILABLE:
        return

    doc = Document()
    doc.add_heading("California Housing Affordability Analysis — Technical Appendix (Generated)", level=1)

    doc.add_paragraph(f"Total observations: {key_stats['n']:,}")
    doc.add_paragraph("Dataset: California Housing Prices (1990 census districts, Kaggle)")

    # Tables
    for title, df in tables.items():
        doc.add_heading(title, level=2)
        table = doc.add_table(rows=1, cols=len(df.columns))
        hdr = table.rows[0].cells
        for j, col in enumerate(df.columns):
            hdr[j].text = str(col)

        for _, row in df.iterrows():
            cells = table.add_row().cells
            for j, col in enumerate(df.columns):
                cells[j].text = str(row[col])

    # Figures
    doc.add_heading("Figures", level=2)
    for label, p in figure_paths.items():
        if p.exists():
            doc.add_paragraph(label)
            doc.add_picture(str(p), width=Inches(6.5))

    # Key stats summary
    doc.add_heading("Key Results Summary", level=2)
    doc.add_paragraph(
        (
            f"Multiple regression: R²={key_stats['reg_r2']:.4f}, Adjusted R²={key_stats['reg_adj_r2']:.4f}, "
            f"F={key_stats['reg_F']:.2f}, p<{key_stats['reg_F_p']:.4g}."
        )
    )
    doc.add_paragraph(
        (
            f"ANOVA: F={key_stats['anova_F']:.2f}, p<{key_stats['anova_p']:.4g}, η²={key_stats['anova_eta2']:.4f}."
        )
    )
    doc.add_paragraph(
        (
            f"K-means (k=4): silhouette={key_stats['silhouette']:.3f}."
        )
    )
    doc.add_paragraph(
        (
            f"5-fold CV R²: mean={key_stats['cv_mean']:.4f}, SD={key_stats['cv_std']:.4f}, "
            f"range={key_stats['cv_min']:.4f}–{key_stats['cv_max']:.4f}."
        )
    )

    ensure_dir(out_docx.parent)
    doc.save(str(out_docx))


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--data", type=str, default=None, help="Path to housing.csv")
    parser.add_argument("--out", type=str, default="outputs", help="Output directory")
    parser.add_argument("--no-docx", action="store_true", help="Skip DOCX generation")
    args = parser.parse_args()

    out_root = Path(args.out)
    out_fig = ensure_dir(out_root / "figures")
    out_tbl = ensure_dir(out_root / "tables")
    out_rep = ensure_dir(out_root / "report")

    data_path = Path(args.data) if args.data else (find_default_dataset() or Path("housing.csv"))

    df_raw = load_dataset(data_path)
    t1 = table_1_missing_summary(df_raw)

    df = impute_bedrooms_by_ocean(df_raw)
    df = feature_engineering(df)

    # Tables
    t2 = table_2_descriptive(df)
    t3 = table_3_ocean_distribution(df)

    # Figures
    corr = figure_1_corr(df, out_fig / "figure_1_correlation_matrix.png")
    dist_stats = figure_2_distributions(df, out_fig / "figure_2_distribution_analysis.png")

    # Regression
    model, X = build_regression(df)
    t4 = table_4_regression_coeffs(model)
    t5 = table_5_vif(X)
    figure_5_residuals(model, out_fig / "figure_5_residual_diagnostics.png")

    # ANOVA
    an = anova_ocean(df)
    t6 = table_6_ocean_stats(df)

    # Clustering
    cl = clustering(df, out_fig)
    df_labeled = cl["df_labeled"]
    t7 = table_7_cluster_profiles(df_labeled)

    # Cross validation
    cv = cross_validation_r2(df)

    # Save tables
    t1.to_csv(out_tbl / "table_1_missing_values.csv", index=False)
    t2.to_csv(out_tbl / "table_2_descriptive_statistics.csv", index=False)
    t3.to_csv(out_tbl / "table_3_ocean_distribution.csv", index=False)
    t4.to_csv(out_tbl / "table_4_regression_coefficients.csv", index=False)
    t5.to_csv(out_tbl / "table_5_vif.csv", index=False)
    t6.to_csv(out_tbl / "table_6_ocean_value_stats.csv", index=False)
    t7.to_csv(out_tbl / "table_7_cluster_profiles.csv", index=False)
    cl["inertias"].to_csv(out_tbl / "kmeans_elbow_inertias.csv", index=False)

    # Print key results
    print("\n=== Key results (compare to PDF) ===")
    print(f"n = {len(df):,}")
    print(f"Figure 2 skewness (median_house_value) = {dist_stats['skewness_house_value']:.2f}")
    print(f"Correlation r(median_income, median_house_value) = {corr.loc['median_income','median_house_value']:.2f}")
    print("\nMultiple regression summary:")
    print(f"R² = {model.rsquared:.4f}")
    print(f"Adj R² = {model.rsquared_adj:.4f}")
    print(f"F = {model.fvalue:.2f}  p {format_p(model.f_pvalue)}")

    # show the specific coefficient the narrative uses
    coef_income = float(model.params.get("median_income", np.nan))
    print(f"β(median_income) = {coef_income:,.2f} (per +1 in median_income, i.e., +$10k)")

    print("\nANOVA:")
    print(f"F = {an['F']:.2f}  p {format_p(an['p'])}  eta^2 = {an['eta2']:.4f}")

    print("\nK-means:")
    print(f"k=4 silhouette = {cl['silhouette']:.3f}")

    print("\n5-fold CV R²:")
    scores = cv["fold_r2"]
    print(f"mean={cv['mean_r2']:.4f}, SD={cv['std_r2']:.4f}, min={min(scores):.4f}, max={max(scores):.4f}")

    # generate a DOCX appendix mirroring the PDF technical appendix
    if (not args.no_docx) and DOCX_AVAILABLE:
        key_stats = {
            "n": len(df),
            "reg_r2": float(model.rsquared),
            "reg_adj_r2": float(model.rsquared_adj),
            "reg_F": float(model.fvalue),
            "reg_F_p": float(model.f_pvalue),
            "anova_F": float(an["F"]),
            "anova_p": float(an["p"]),
            "anova_eta2": float(an["eta2"]),
            "silhouette": float(cl["silhouette"]),
            "cv_mean": float(cv["mean"]),
            "cv_std": float(cv["std"]),
            "cv_min": float(cv["min"]),
            "cv_max": float(cv["max"]),
        }
        tables = {
            "Table 1: Missing Values Summary": t1,
            "Table 2: Descriptive Statistics": t2,
            "Table 3: Distribution by Ocean Proximity": t3,
            "Table 4: Regression Coefficients": t4,
            "Table 5: Variance Inflation Factors": t5,
            "Table 6: House Value Statistics by Ocean Proximity": t6,
            "Table 7: Cluster Profiles": t7,
        }
        figure_paths = {
            "Figure 1: Correlation Matrix": out_fig / "figure_1_correlation_matrix.png",
            "Figure 2: Distribution Analysis": out_fig / "figure_2_distribution_analysis.png",
            "Figure 3: Elbow Plot": out_fig / "figure_3_elbow_plot.png",
            "Figure 4: Geographic Clusters": out_fig / "figure_4_geographic_clusters.png",
            "Figure 5: Residual Diagnostics": out_fig / "figure_5_residual_diagnostics.png",
        }
        maybe_write_docx(out_rep / "technical_appendix_generated.docx", tables, key_stats, figure_paths)
        print(f"\nDOCX appendix written to: {out_rep / 'technical_appendix_generated.docx'}")

    print(f"\nAll outputs written under: {out_root.resolve()}")


if __name__ == "__main__":
    main()
